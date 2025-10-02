import os
import logging
import io
from datetime import datetime
from telegram import Update, ReplyKeyboardMarkup, ReplyKeyboardRemove
from telegram.ext import Application, CommandHandler, MessageHandler, ContextTypes, ConversationHandler, filters
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from dotenv import load_dotenv
from http.server import HTTPServer, BaseHTTPRequestHandler
from threading import Thread
import time

# ==================== ПРОСТОЙ ВЕБ-СЕРВЕР ДЛЯ RENDER ====================
class HealthCheckHandler(BaseHTTPRequestHandler):
    def do_GET(self):
        if self.path == '/':
            self.send_response(200)
            self.send_header('Content-type', 'text/html')
            self.end_headers()
            self.wfile.write(b"""
                <html>
                    <head><title>Insurance Bot</title></head>
                    <body>
                        <h1>🤖 Бот страхования работает!</h1>
                        <p>Insurance Bot is ONLINE and ready to receive applications.</p>
                        <p>🕒 Статус: <strong>Активен</strong></p>
                        <p>📅 Время сервера: """ + datetime.now().strftime('%Y-%m-%d %H:%M:%S').encode() + b"""</p>
                    </body>
                </html>
            """)
        else:
            self.send_response(404)
            self.end_headers()

def run_health_check():
    port = int(os.environ.get('PORT', 10000))
    server = HTTPServer(('0.0.0.0', port), HealthCheckHandler)
    print(f"✅ Веб-сервер запущен на порту {port}")
    server.serve_forever()

# Запускаем веб-сервер в фоне
health_thread = Thread(target=run_health_check, daemon=True)
health_thread.start()

# ==================== ЗАГРУЗКА ПЕРЕМЕННЫХ ====================
load_dotenv()

# ==================== НАСТРОЙКА ЛОГИРОВАНИЯ ====================
logging.basicConfig(
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    level=logging.INFO
)

print("🚀 Начинается запуск Telegram бота...")
print("=== 🔥 ОБНОВЛЕННАЯ ВЕРСИЯ 2.0 ===")
print("=== ✅ НАВИГАЦИЯ + ВОДИТЕЛЬСКИЕ ПРАВА ===")

# ==================== СОСТОЯНИЯ РАЗГОВОРА ====================
(
    START, CHOOSE_OWNER_INSURER, INSURER_FIO, INSURER_BIRTHDATE, INSURER_PASSPORT_SERIES_NUMBER,
    INSURER_PASSPORT_ISSUE_DATE, INSURER_PASSPORT_ISSUED_BY, INSURER_PASSPORT_DEPARTMENT_CODE,
    INSURER_REGISTRATION, OWNER_FIO, OWNER_BIRTHDATE, OWNER_PASSPORT_SERIES_NUMBER,
    OWNER_PASSPORT_ISSUE_DATE, OWNER_PASSPORT_ISSUED_BY, OWNER_PASSPORT_DEPARTMENT_CODE,
    INSURER_LICENSE, INSURER_LICENSE_ISSUE_DATE, INSURER_LICENSE_EXPIRY, VEHICLE_BRAND,
    VEHICLE_MODEL, VEHICLE_YEAR, VEHICLE_POWER, VEHICLE_REG_NUMBER, VEHICLE_VIN,
    VEHICLE_DOC_TYPE, VEHICLE_DOC_DETAILS, VEHICLE_DOC_ISSUE_DATE, DRIVERS_CHOICE,
    DRIVER_FIO, DRIVER_LICENSE, DRIVER_LICENSE_ISSUE_DATE, DRIVER_LICENSE_EXPIRY, INSURER_PHONE,
    CONFIRMATION
) = range(34)

user_data = {}

class WordGenerator:
    @staticmethod
    def generate_application_docx(data):
        """Генерация Word документа с заявкой"""
        doc = Document()
        
        # Заголовок
        title = doc.add_heading('ЗАЯВКА НА СТРАХОВАНИЕ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Дата
        date_paragraph = doc.add_paragraph()
        date_paragraph.add_run(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}").bold = True
        doc.add_paragraph()
        
        # Раздел: Страхователь
        doc.add_heading('СТРАХОВАТЕЛЬ', level=1)
        
        insurer_info = [
            f"ФИО: {data.get('insurer_fio', 'Не указано')}",
            f"Дата рождения: {data.get('insurer_birthdate', 'Не указано')}",
            f"Паспорт: {data.get('insurer_passport_series_number', 'Не указано')}",
            f"Дата выдачи паспорта: {data.get('insurer_passport_issue_date', 'Не указано')}",
            f"Кем выдан: {data.get('insurer_passport_issued_by', 'Не указано')}",
            f"Код подразделения: {data.get('insurer_passport_department_code', 'Не указано')}",
            f"Прописка: {data.get('insurer_registration', 'Не указано')}"
        ]
        
        for info in insurer_info:
            doc.add_paragraph(info)
        
        doc.add_paragraph()
        
        # Раздел: Собственник
        doc.add_heading('СОБСТВЕННИК', level=1)
        
        if not data.get('is_same_person', True):
            owner_info = [
                f"ФИО: {data.get('owner_fio', 'Не указано')}",
                f"Дата рождения: {data.get('owner_birthdate', 'Не указано')}",
                f"Паспорт: {data.get('owner_passport_series_number', 'Не указано')}",
                f"Дата выдачи паспорта: {data.get('owner_passport_issue_date', 'Не указано')}",
                f"Кем выдан: {data.get('owner_passport_issued_by', 'Не указано')}",
                f"Код подразделения: {data.get('owner_passport_department_code', 'Не указано')}"
            ]
            
            for info in owner_info:
                doc.add_paragraph(info)
        else:
            doc.add_paragraph("Собственник и страхователь - одно лицо")
        
        doc.add_paragraph()
        
        # Водительское удостоверение страхователя
        doc.add_heading('ВОДИТЕЛЬСКОЕ УДОСТОВЕРЕНИЕ СТРАХОВАТЕЛЯ', level=1)
        
        license_info = [
            f"В/у: {data.get('insurer_license', 'Не указано')}",
            f"Дата выдачи: {data.get('insurer_license_issue_date', 'Не указано')}",
            f"Срок действия: {data.get('insurer_license_expiry', 'Не указано')}"
        ]
        
        for info in license_info:
            doc.add_paragraph(info)
        
        doc.add_paragraph()
        
        # Раздел: Транспортное средство
        doc.add_heading('ТРАНСПОРТНОЕ СРЕДСТВО', level=1)
        
        vehicle_info = [
            f"Марка: {data.get('vehicle_brand', 'Не указано')}",
            f"Модель: {data.get('vehicle_model', 'Не указано')}",
            f"Год выпуска: {data.get('vehicle_year', 'Не указано')}",
            f"Мощность: {data.get('vehicle_power', 'Не указано')} л.с.",
            f"Госномер: {data.get('vehicle_reg_number', 'Не указано')}",
            f"VIN: {data.get('vehicle_vin', 'Не указано')}",
            f"Документ: {data.get('vehicle_doc_type', 'Не указано')} {data.get('vehicle_doc_details', 'Не указано')}",
            f"Дата выдачи документа: {data.get('vehicle_doc_issue_date', 'Не указано')}"
        ]
        
        for info in vehicle_info:
            doc.add_paragraph(info)
        
        doc.add_paragraph()
        
        # Раздел: Водители
        doc.add_heading('ВОДИТЕЛИ', level=1)
        
        drivers = data.get('drivers', [])
        if drivers:
            for i, driver in enumerate(drivers, 1):
                driver_paragraph = doc.add_paragraph()
                driver_paragraph.add_run(f'Водитель {i}: ').bold = True
                driver_paragraph.add_run(f"{driver.get('fio', 'Не указано')}")
                
                doc.add_paragraph(f"   В/у: {driver.get('license', 'Не указано')}")
                doc.add_paragraph(f"   Дата выдачи: {driver.get('license_issue_date', 'Не указано')}")
                doc.add_paragraph(f"   Срок действия: {driver.get('license_expiry', 'Не указано')}")
                doc.add_paragraph()
        else:
            doc.add_paragraph("Водители не указаны")
        
        # Телефон
        doc.add_paragraph()
        phone_paragraph = doc.add_paragraph()
        phone_paragraph.add_run("Телефон для связи: ").bold = True
        phone_paragraph.add_run(f"{data.get('insurer_phone', 'Не указан')}")
        
        # Подпись
        doc.add_paragraph()
        doc.add_paragraph("Заявка успешно оформлена!").bold = True
        doc.add_paragraph("В течении 1 часа с Вами свяжется менеджер, для возможного уточнения деталей и дальнейшего оформления!")
        doc.add_paragraph("С Уважением, АО 'Альфастрахование'").bold = True
        
        return doc

def get_navigation_keyboard():
    """Клавиатура для навигации"""
    return ReplyKeyboardMarkup([
        ["⬅️ Назад", "🏠 В начало"]
    ], resize_keyboard=True)

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Начало разговора"""
    user = update.message.from_user
    await update.message.reply_text(
        f"Добро пожаловать, {user.first_name}!\n"
        "Я помогу собрать информацию для страховки.\n\n"
        "Собственник и страхователь - одно лицо?",
        reply_markup=ReplyKeyboardMarkup([
            ["✅ Одно лицо", "❌ Разные лица"]
        ], one_time_keyboard=True, resize_keyboard=True)
    )
    return CHOOSE_OWNER_INSURER

async def choose_owner_insurer(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка выбора типа собственника/страхователя"""
    if update.message.text in ["⬅️ Назад", "🏠 В начало"]:
        return await start(update, context)
    
    choice = update.message.text
    user_id = update.message.from_user.id
    
    user_data[user_id] = {
        'is_same_person': choice == "✅ Одно лицо",
        'drivers': []
    }
    
    await update.message.reply_text(
        "Введите ФИО страхователя полностью:",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_FIO

async def insurer_fio(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Получение ФИО страхователя"""
    if update.message.text in ["⬅️ Назад", "🏠 В начало"]:
        return await start(update, context)
    
    user_id = update.message.from_user.id
    user_data[user_id]['insurer_fio'] = update.message.text
    
    await update.message.reply_text(
        "Введите дату рождения страхователя (в формате ДД.ММ.ГГГГ):\n"
        "Пример: 15.05.1990",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_BIRTHDATE

async def insurer_birthdate(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Получение даты рождения страхователя"""
    if update.message.text == "⬅️ Назад":
        await update.message.reply_text(
            "Введите ФИО страхователя полностью:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_FIO
    elif update.message.text == "🏠 В начало":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    try:
        datetime.strptime(update.message.text, '%d.%m.%Y')
        user_data[user_id]['insurer_birthdate'] = update.message.text
    except ValueError:
        await update.message.reply_text(
            "Неверный формат даты. Введите в формате ДД.ММ.ГГГГ:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_BIRTHDATE
    
    await update.message.reply_text(
        "Введите серию и номер паспорта страхователя:\n"
        "Пример: 1234 567890",
        reply_markup=get_navigation_keyboard()
    )
    return INSURER_PASSPORT_SERIES_NUMBER

# ... (пропущены промежуточные обработчики для экономии места)
# Все остальные обработчики остаются без изменений

async def insurer_phone(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Получение телефона для связи"""
    if update.message.text == "⬅️ Назад":
        await update.message.reply_text(
            "Выберите действие с водителями:",
            reply_markup=ReplyKeyboardMarkup([
                ["📋 Скопировать страхователя", "👤 Добавить водителя"],
                ["✅ Завершить добавление"],
                ["⬅️ Назад", "🏠 В начало"]
            ], resize_keyboard=True)
        )
        return DRIVERS_CHOICE
    elif update.message.text == "🏠 В начало":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    if user_id not in user_data:
        await update.message.reply_text("Пожалуйста, начните с команды /start")
        return ConversationHandler.END
        
    user_data[user_id]['insurer_phone'] = update.message.text
    
    # Переход к подтверждению
    await update.message.reply_text(
        "✅ Все данные собраны!\n\n"
        "Нажмите кнопку ниже для подтверждения и отправки заявки:",
        reply_markup=ReplyKeyboardMarkup([
            ["✅ Подтвердить и отправить"],
            ["⬅️ Назад", "🏠 В начало"]
        ], resize_keyboard=True)
    )
    return CONFIRMATION

async def confirmation_handler(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Обработка подтверждения заявки"""
    if update.message.text == "⬅️ Назад":
        await update.message.reply_text(
            "Введите телефон для связи:",
            reply_markup=get_navigation_keyboard()
        )
        return INSURER_PHONE
    elif update.message.text == "🏠 В начало":
        return await start(update, context)
    
    user_id = update.message.from_user.id
    if user_id not in user_data:
        await update.message.reply_text("Пожалуйста, начните с команды /start")
        return ConversationHandler.END
        
    # Вызываем функцию подтверждения и отправки
    return await send_confirmation(update, context)

async def send_confirmation(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Подтверждение и отправка данных"""
    user_id = update.message.from_user.id
    if user_id not in user_data:
        await update.message.reply_text("Пожалуйста, начните с команды /start")
        return ConversationHandler.END
        
    data = user_data[user_id]
    
    try:
        # Формируем детальное сообщение для Telegram
        manager_message = "🚗 СРОЧНАЯ ЗАЯВКА НА СТРАХОВАНИЕ\n\n"
        
        manager_message += "👤 СТРАХОВАТЕЛЬ:\n"
        manager_message += f"ФИО: {data.get('insurer_fio', 'Не указано')}\n"
        manager_message += f"Дата рождения: {data.get('insurer_birthdate', 'Не указано')}\n"
        manager_message += f"Паспорт: {data.get('insurer_passport_series_number', 'Не указано')}\n"
        manager_message += f"Дата выдачи: {data.get('insurer_passport_issue_date', 'Не указано')}\n"
        manager_message += f"Кем выдан: {data.get('insurer_passport_issued_by', 'Не указано')}\n"
        manager_message += f"Код подразделения: {data.get('insurer_passport_department_code', 'Не указано')}\n"
        manager_message += f"Прописка: {data.get('insurer_registration', 'Не указано')}\n\n"
        
        # ДОБАВЛЕНО: Водительское удостоверение страхователя
        manager_message += "🚗 ВОДИТЕЛЬСКОЕ УДОСТОВЕРЕНИЕ СТРАХОВАТЕЛЯ:\n"
        manager_message += f"Номер: {data.get('insurer_license', 'Не указано')}\n"
        manager_message += f"Дата выдачи: {data.get('insurer_license_issue_date', 'Не указано')}\n"
        manager_message += f"Срок действия: {data.get('insurer_license_expiry', 'Не указано')}\n\n"
        
        if not data.get('is_same_person', True):
            manager_message += "👤 СОБСТВЕННИК:\n"
            manager_message += f"ФИО: {data.get('owner_fio', 'Не указано')}\n"
            manager_message += f"Дата рождения: {data.get('owner_birthdate', 'Не указано')}\n"
            manager_message += f"Паспорт: {data.get('owner_passport_series_number', 'Не указано')}\n"
            manager_message += f"Дата выдачи: {data.get('owner_passport_issue_date', 'Не указано')}\n"
            manager_message += f"Кем выдан: {data.get('owner_passport_issued_by', 'Не указано')}\n"
            manager_message += f"Код подразделения: {data.get('owner_passport_department_code', 'Не указано')}\n\n"
        else:
            manager_message += "👤 СОБСТВЕННИК:\n"
            manager_message += "Собственник и страхователь - одно лицо\n\n"
        
        manager_message += "🚗 ТРАНСПОРТНОЕ СРЕДСТВО:\n"
        manager_message += f"Марка: {data.get('vehicle_brand', 'Не указано')}\n"
        manager_message += f"Модель: {data.get('vehicle_model', 'Не указано')}\n"
        manager_message += f"Год выпуска: {data.get('vehicle_year', 'Не указано')}\n"
        manager_message += f"Мощность: {data.get('vehicle_power', 'Не указано')} л.с.\n"
        manager_message += f"Госномер: {data.get('vehicle_reg_number', 'Не указано')}\n"
        manager_message += f"VIN: {data.get('vehicle_vin', 'Не указано')}\n"
        manager_message += f"Документ: {data.get('vehicle_doc_type', 'Не указано')} {data.get('vehicle_doc_details', 'Не указано')}\n"
        manager_message += f"Дата выдачи: {data.get('vehicle_doc_issue_date', 'Не указано')}\n\n"
        
        manager_message += "👥 ВОДИТЕЛИ:\n"
        drivers = data.get('drivers', [])
        if drivers:
            for i, driver in enumerate(drivers, 1):
                manager_message += f"{i}. {driver.get('fio', 'Не указано')}\n"
                manager_message += f"   В/у: {driver.get('license', 'Не указано')}\n"
                manager_message += f"   Дата выдачи: {driver.get('license_issue_date', 'Не указано')}\n"
                manager_message += f"   Срок действия: {driver.get('license_expiry', 'Не указано')}\n\n"
        else:
            manager_message += "Водители не указаны\n\n"
        
        manager_message += f"📞 Телефон: {data.get('insurer_phone', 'Не указан')}\n"
        manager_message += f"📅 Дата заявки: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
        
        # Отправляем детальное уведомление менеджеру в Telegram
        MANAGER_CHAT_ID = os.getenv('MANAGER_CHAT_ID')
        if MANAGER_CHAT_ID:
            try:
                if len(manager_message) > 4096:
                    parts = [manager_message[i:i+4096] for i in range(0, len(manager_message), 4096)]
                    for part in parts:
                        await context.bot.send_message(chat_id=int(MANAGER_CHAT_ID), text=part)
                else:
                    await context.bot.send_message(chat_id=int(MANAGER_CHAT_ID), text=manager_message)
                
                print(f"✅ Текстовое уведомление отправлено менеджеру {MANAGER_CHAT_ID}")
            except Exception as e:
                print(f"❌ Ошибка отправки в Telegram: {e}")
        
        # Создаем Word документ
        doc = WordGenerator.generate_application_docx(data)
        
        # Сохраняем Word в байты
        file_stream = io.BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        file_stream.name = f"Заявка_{data.get('insurer_fio', 'Клиент')}_{datetime.now().strftime('%d%m%Y_%H%M')}.docx"
        
        # Отправляем Word документ менеджеру
        if MANAGER_CHAT_ID:
            try:
                await context.bot.send_document(
                    chat_id=int(MANAGER_CHAT_ID),
                    document=file_stream,
                    caption=f"📄 Заявка от {data.get('insurer_fio', 'Клиент')}"
                )
                print(f"✅ Word документ отправлен менеджеру {MANAGER_CHAT_ID}")
            except Exception as e:
                print(f"❌ Ошибка отправки Word менеджеру: {e}")
        
        # Отправляем подтверждение клиенту
        await update.message.reply_text(
            "✅ Заявка успешно оформлена!\n\n"
            "В течении 1 часа с Вами свяжется менеджер, для возможного уточнения деталей и дальнейшего оформления!\n\n"
            "С Уважением, АО 'Альфастрахование'",
            reply_markup=ReplyKeyboardRemove()
        )
        
        # Отправляем текстовую копию клиенту
        client_message = "📋 Ваша заявка:\n\n" + manager_message
        if len(client_message) > 4096:
            parts = [client_message[i:i+4096] for i in range(0, len(client_message), 4096)]
            for part in parts:
                await update.message.reply_text(part)
        else:
            await update.message.reply_text(client_message)
        
        # Отправляем Word документ клиенту
        file_stream.seek(0)
        await update.message.reply_document(
            document=file_stream,
            caption="📄 Ваша заявка на страхование"
        )
        
    except Exception as e:
        print(f"❌ Критическая ошибка: {e}")
        await update.message.reply_text(
            "Произошла непредвиденная ошибка. "
            "Пожалуйста, попробуйте позже.",
            reply_markup=ReplyKeyboardRemove()
        )
    
    # Очищаем данные пользователя
    if user_id in user_data:
        del user_data[user_id]
    
    return ConversationHandler.END

async def cancel(update: Update, context: ContextTypes.DEFAULT_TYPE) -> int:
    """Отмена разговора"""
    user_id = update.message.from_user.id
    if user_id in user_data:
        del user_data[user_id]
    
    await update.message.reply_text(
        "Заявка отменена.",
        reply_markup=ReplyKeyboardRemove()
    )
    return ConversationHandler.END

def main():
    """Запуск бота"""
    TOKEN = os.getenv('TELEGRAM_BOT_TOKEN')
    
    if not TOKEN:
        logging.error("❌ Ошибка: не задан TELEGRAM_BOT_TOKEN")
        return
    
    try:
        application = Application.builder().token(TOKEN).build()
        
        conv_handler = ConversationHandler(
            entry_points=[CommandHandler('start', start)],
            states={
                CHOOSE_OWNER_INSURER: [MessageHandler(filters.TEXT & ~filters.COMMAND, choose_owner_insurer)],
                INSURER_FIO: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_fio)],
                INSURER_BIRTHDATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_birthdate)],
                INSURER_PASSPORT_SERIES_NUMBER: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_passport_series_number)],
                INSURER_PASSPORT_ISSUE_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_passport_issue_date)],
                INSURER_PASSPORT_ISSUED_BY: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_passport_issued_by)],
                INSURER_PASSPORT_DEPARTMENT_CODE: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_passport_department_code)],
                INSURER_REGISTRATION: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_registration)],
                OWNER_FIO: [MessageHandler(filters.TEXT & ~filters.COMMAND, owner_fio)],
                OWNER_BIRTHDATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, owner_birthdate)],
                OWNER_PASSPORT_SERIES_NUMBER: [MessageHandler(filters.TEXT & ~filters.COMMAND, owner_passport_series_number)],
                OWNER_PASSPORT_ISSUE_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, owner_passport_issue_date)],
                OWNER_PASSPORT_ISSUED_BY: [MessageHandler(filters.TEXT & ~filters.COMMAND, owner_passport_issued_by)],
                OWNER_PASSPORT_DEPARTMENT_CODE: [MessageHandler(filters.TEXT & ~filters.COMMAND, owner_passport_department_code)],
                INSURER_LICENSE: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_license)],
                INSURER_LICENSE_ISSUE_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_license_issue_date)],
                INSURER_LICENSE_EXPIRY: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_license_expiry)],
                VEHICLE_BRAND: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_brand)],
                VEHICLE_MODEL: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_model)],
                VEHICLE_YEAR: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_year)],
                VEHICLE_POWER: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_power)],
                VEHICLE_REG_NUMBER: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_reg_number)],
                VEHICLE_VIN: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_vin)],
                VEHICLE_DOC_TYPE: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_doc_type)],
                VEHICLE_DOC_DETAILS: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_doc_details)],
                VEHICLE_DOC_ISSUE_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, vehicle_doc_issue_date)],
                DRIVERS_CHOICE: [MessageHandler(filters.TEXT & ~filters.COMMAND, drivers_choice)],
                DRIVER_FIO: [MessageHandler(filters.TEXT & ~filters.COMMAND, driver_fio)],
                DRIVER_LICENSE: [MessageHandler(filters.TEXT & ~filters.COMMAND, driver_license)],
                DRIVER_LICENSE_ISSUE_DATE: [MessageHandler(filters.TEXT & ~filters.COMMAND, driver_license_issue_date)],
                DRIVER_LICENSE_EXPIRY: [MessageHandler(filters.TEXT & ~filters.COMMAND, driver_license_expiry)],
                INSURER_PHONE: [MessageHandler(filters.TEXT & ~filters.COMMAND, insurer_phone)],
                CONFIRMATION: [MessageHandler(filters.TEXT & ~filters.COMMAND, confirmation_handler)],
            },
            fallbacks=[CommandHandler('cancel', cancel)]
        )
        
        application.add_handler(conv_handler)
        
        logging.info("🤖 Бот запускается...")
        print("=== БОТ ЗАПУЩЕН НА RENDER ===")
        
        application.run_polling(
            drop_pending_updates=True,
            allowed_updates=Update.ALL_TYPES,
            close_loop=False
        )
        
    except Exception as e:
        logging.error(f"❌ Критическая ошибка: {e}")
        print("Бот остановлен из-за ошибки:", e)
        time.sleep(10)
        main()

if __name__ == '__main__':
    main()
